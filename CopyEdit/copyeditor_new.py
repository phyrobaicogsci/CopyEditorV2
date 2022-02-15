# import the Document class
# from the docx module
from docx.api import Document
import pandas as pd
import re
import requests
API_URL = "https://api-inference.huggingface.co/models/vennify/t5-base-grammar-correction"
headers = {"Authorization": "Bearer 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"}

# nlp = spacy.load('en_core_web_sm')
def start():
    # create an instance of a
    # word document we want to open
    doc = Document(r'input.docx')
    document = Document('input.docx')
    table = document.tables[:]

    data = []

    keys = None
    for t in table:
        for i, row in enumerate(t.rows):
            text = (cell.text for cell in row.cells)
            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)

    df = pd.DataFrame(data)


    for i in data:
        for j in i.keys():
            doc.add_paragraph("\n")
            doc.add_paragraph(i[j])
        
    for i in data:
        for key,value in i.items():
            doc.add_paragraph("\n")
            doc.add_paragraph(key)
        


    # for printing the complete document

        


    document = Document()
        
    doc.save('demo.docx')

    # Rules implemented here 
                    
    def docx_replace_regex(doc, regex , replace): 
        for p in doc.paragraphs: 
            if regex.search(p.text): 
                inline = p.runs 
                # Loop added to work with runs (strings with same style) 
                for i in range(len(inline)): 
                    if regex.search(inline[i].text): 
                        text = regex.sub(replace, inline[i].text) 
                        inline[i].text = text 
    
        for table in doc.tables: 
            for row in table.rows: 
                for cell in row.cells: 
                    docx_replace_regex(cell, regex , replace) 


    regex1 = re.compile("") 
    replace1 = r"" 

    filename = "test.docx" 
    doc = Document('demo.docx') 
    docx_replace_regex(doc, regex1 , replace1)
    doc.save('demo.docx')
    print('\n','\n','file overwritten')

    #COPYEDIT SUGGESTIONS
    import language_tool_python
    tool = language_tool_python.LanguageTool('en-US')
    
    doc1 = Document(r'demo.docx')

    def getText(filename):
        fullText = []
        for para in doc1.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    # get the matches
    matches = tool.check(getText('demo'))
    corrected_text = getText('demo')

    corrected_text = tool.correct(getText('demo'))

    corrected_text = re.sub("(?<=[\+\-\*\=\×])(?=[^\s])|(?<=[^\s])(?=[\+\-\*\=\×])",r" ",corrected_text)
    
    corrected_text = re.sub("(?<=[ \/])(?=[^\s+])|(?<=[^\s+])(?=[\/ ])",r"",corrected_text)
    
    corrected_text = re.sub("(?<=[a-z])\s(?=[:])",r"",corrected_text)
    corrected_text = re.sub("(?<=[:])\s(?=[a-z])",r"",corrected_text)
    
    corrected_text = re.sub("(?<=(\d:))(?=[^\s])",r" ",corrected_text)

    corrected_text = re.sub("\d :",lambda m: m.group(0).replace(" ",''),corrected_text)

    corrected_text = re.sub("(?<=[:])(?=[^\s])(?=[A-Z])",r" ",corrected_text)
    corrected_text = re.sub("\d:[a-z]",lambda m: m.group(0).upper(),corrected_text)
    corrected_text = re.sub("\d: [a-z]",lambda m: m.group(0).upper(),corrected_text)
    corrected_text = re.sub("[A-Z]:",lambda m: m.group(0).lower(),corrected_text)
   
    corrected_text = re.sub("[A-Z]\s(?=[\+\-\=\×\/])",lambda m: m.group(0).lower(),corrected_text)
   
    corrected_text = re.sub("(?<=[\.\:\=\(\]\[\)\{\}])(?=[\.\:\=\(\]\[\)\{\}])",r" ",corrected_text)

    corrected_text = re.sub("(?<=[\d])(?=[aA-zZ])",r" ",corrected_text)

    corrected_text = re.sub("(?<=[aA-zZ])(?=[\d])",r" ",corrected_text)
    
    
    corrected_text = corrected_text.replace("???", "?")
    corrected_text = corrected_text.replace("??", "?")
    corrected_text = corrected_text.replace("? ? ?", "?")
    corrected_text = corrected_text.replace("? ?", "?")



    corrected_text = corrected_text.replace("!!", "!")
    corrected_text = corrected_text.replace("!!!", "!")
    corrected_text = corrected_text.replace("! !", "!")
    corrected_text = corrected_text.replace("! ! !", "!")

    corrected_text = corrected_text.replace("kilometre", "kilometer")
    corrected_text = corrected_text.replace("Kilometre", "Kilometer")
    corrected_text = corrected_text.replace("KILOMETRE", "KILOMETER")

    corrected_text = corrected_text.replace("kilometres", "kilometers")
    corrected_text = corrected_text.replace("Kilometres", "Kilometers")
    corrected_text = corrected_text.replace("KILOMETRES", "KILOMETERS")

    corrected_text = corrected_text.replace("and ,", "and")
    corrected_text = corrected_text.replace("and,", "and")

    corrected_text = re.sub("(?<=[,])(?=[aA-zZ])",r" ",corrected_text)
    
    corrected_text = re.sub("(?<=[aA-zZ])[\s](?=[,])",r"",corrected_text)
    
    corrected_text = re.sub("(?<=[\w])[\s](?=[a][n][d])",r", ",corrected_text)
    
    corrected_text = re.sub(r"\b(one|two|three|four|five|six|seven|eight|nine|ten)\b \b(half|third|forth|fifth|sixth|seventh|eighth|ninth|tenth)\b",lambda m: m.group(0).replace(" ","-"),corrected_text)
    
    corrected_text = corrected_text.replace("kms/hr", "km/h")
    corrected_text = corrected_text.replace("Kms/hr", "km/h")
    corrected_text = corrected_text.replace("KMS/HR", "km/h")

    corrected_text = corrected_text.replace("kms/Hr", "km/h")
    corrected_text = corrected_text.replace("Kms/Hr", "km/h")

    corrected_text = corrected_text.replace("kms/hrs", "km/h")
    corrected_text = corrected_text.replace("Kms/hrs", "km/h")
    corrected_text = corrected_text.replace("KMS/HRS", "km/h")
    corrected_text = corrected_text.replace("KMs/Hrs", "km/h")

    corrected_text = corrected_text.replace("km/HS", "km/h")
    corrected_text = corrected_text.replace("Kms/Hrs", "km/h")

    corrected_text = re.sub("¢\d+",lambda m: m.group(0).split("¢")[1]+"¢" ,corrected_text)
    
    corrected_text = re.sub("¢\s\d+",lambda m: m.group(0).split(" ")[1]+"¢" ,corrected_text)
    
    corrected_text = corrected_text.replace("Litre", "liters")
    corrected_text = corrected_text.replace("litre", "liters")
    corrected_text = corrected_text.replace(" 1 liters", " 1 liter")
    corrected_text = corrected_text.replace("Liter", "liter")
    corrected_text = corrected_text.replace(" l.", " L.")
    corrected_text = corrected_text.replace(" l?", " L?")
    corrected_text = corrected_text.replace(" l,", " L,")
    corrected_text = corrected_text.replace("ml", "mL")
    corrected_text = corrected_text.replace("millilitre", "milliliter")
    corrected_text = corrected_text.replace("Millilitre", "Milliliter")
    orrected_text = corrected_text.replace("millilitres", "milliliters")
    corrected_text = corrected_text.replace("Millilitres", "Milliliters")
    corrected_text = corrected_text.replace("literss", "liters")
    corrected_text = corrected_text.replace(" 0 liters", " 0 liter")

    doc1 = Document()
    doc1.add_paragraph(corrected_text)
    doc1.save('corrected.docx')
    
    a=corrected_text.split("\n")
    while("" in a) :
        a.remove("")
    suggestions =[]
    grammar = []
    for i in a:
        if i[-1] != "." :
            if i[-1] != ",":
                if i[-1] != "?":
                    if i[-1] != "!":
                        if " " in i:
                            if "+" not in i:
                                if "÷" not in i:
                                    if "×" not in i:
                                        if "-" not in i:
                                            if "/" not in i:
                                                if "=" not in i:
                                                    if "When" in i or "What" in i or "Why" in i or "How" in i or "Where" in i or "Which" in i:
                                                        suggestions.append(f"<span>Question mark missing after:</span> {i} ")
                                                    else:
                                                        suggestions.append(f"<span>Full stop missing after: </span>{i} ")


    
    
    
    

        def query(payload):
            response = requests.post(API_URL, headers=headers, json=payload)
            return response.json()

        if not any(word in i for word in ["+","-","*","x","=","/"]) and not len(i)==1:
            try:
                output = query({"inputs": i})[0]['generated_text']
                print("original",i)
                print("edited",output)
                
            except:
                output = query({"inputs": i})
              
            if i!=output:
                grammar.append(f"<span>Grammar suggestion:</span> {output}")


    print("Done!")
    return corrected_text,suggestions,grammar